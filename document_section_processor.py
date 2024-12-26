import logging
from io import BytesIO

import matplotlib.pyplot as plt
import ollama
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from bangla_to_unicode import BanglaToUnicode

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentSectionProcessor:
    CUSTOM_MODEL = "rokomari_bot"

    def __init__(self, modelfile_location: str = "static/Modelfile"):
        self.btu = BanglaToUnicode()
        available_models = [model.model for model in ollama.list().models]

        if self.CUSTOM_MODEL + ":latest" not in available_models:
            try:
                modelfile = open(modelfile_location, "r").read()
                ollama.create(self.CUSTOM_MODEL, modelfile=modelfile)
                logger.info(f"Model {self.CUSTOM_MODEL} created successfully")
            except Exception as e:
                logger.error(e)

    def pipeline(self, doc: Document):
        # the calling order is important
        for section in doc.sections:
            self.adjust_page_layout(section)
            self.remove_headers_footers(section)

        for paragraph in doc.paragraphs:
            self.enforce_content_restrictions(paragraph)
        # applying on whole document
        self.process_tables(doc)
        # self.formate_image(doc)

    # Function to adjust page layout
    def adjust_page_layout(self, section: Section):

        cols = section._sectPr.xpath("./w:cols")[0]
        cols.set(qn("w:num"), "1")
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        page_borders = section._sectPr.xpath("./w:pgBorders")
        if page_borders:
            for border in page_borders:
                section._sectPr.remove(border)

    # Function to process headers and footers
    def remove_headers_footers(self, section: Section):
        header = section.header
        footer = section.footer
        header.is_linked_to_previous = False
        footer.is_linked_to_previous = False
        for paragraph in header.paragraphs:
            paragraph.clear()
        for paragraph in footer.paragraphs:
            paragraph.clear()

    # Function to remove unwanted content
    def enforce_content_restrictions(self, paragraph: Paragraph):
        text = paragraph.text.strip()
        if not text:
            return
        text = self.btu.to_sutonnymj(text)
        response = ollama.generate(model=self.CUSTOM_MODEL, prompt=text)
        refined_text = response.response

        paragraph.text = refined_text
        logger.info(f"{text=}, {refined_text=}")

    def process_tables(self, doc: Document, max_width: Inches = Inches(3)):
        # process this at last as this will work on whole table
        for table in doc.tables:
            table_width = sum(cell.width for cell in table.rows[0].cells)
            logger.info(f"{table_width=}")
            if table_width <= max_width:
                continue

            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            image_stream = self.array_to_image(table_data)

            parent = table._element.getparent()
            table_index = list(parent).index(table._element)
            parent.remove(table._element)
            new_paragraph_element = OxmlElement("w:p")
            parent.insert(table_index, new_paragraph_element)
            paragraph = doc.paragraphs[table_index]

            # Add the image to the new paragraph
            run = paragraph.add_run()
            run.add_picture(image_stream)

    def array_to_image(self, table: list):
        df = pd.DataFrame(table[1:], columns=table[0])
        fig, ax = plt.subplots(figsize=(6, 2))
        ax.axis("tight")
        ax.axis("off")
        table = ax.table(
            cellText=df.values, colLabels=df.columns, loc="center", cellLoc="center"
        )
        image_buffer = BytesIO()
        fig.savefig(image_buffer, format="png", bbox_inches="tight")
        image_buffer.seek(0)
        return image_buffer

    def formate_image(self, doc: Document):
        for shape in doc.inline_shapes:
            rel_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = doc.part.related_parts[rel_id]

            original_image_bytes = image_part.blob
            new_image_bytes = self.__ushape_image(original_image_bytes)
            image_part.blob = new_image_bytes

    def __ushape_image(self, image_bytes: bytes):
        pass
